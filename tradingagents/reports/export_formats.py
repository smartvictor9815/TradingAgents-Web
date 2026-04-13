"""Markdown → DOCX / PDF bytes for API download endpoints."""

from __future__ import annotations

import html
import io
import logging
import re
from pathlib import Path
from typing import Optional
import xml.etree.ElementTree as ET

_log = logging.getLogger("uvicorn.error")


def markdown_to_docx_bytes(md: str) -> bytes:
    """Render Markdown to structured DOCX (headings, lists, tables, code)."""
    try:
        from docx import Document
        from docx.shared import Pt
    except ImportError as e:
        raise RuntimeError("python-docx is not installed; pip install python-docx") from e

    doc = Document()
    normal_style = doc.styles["Normal"]
    normal_style.font.name = "Calibri"
    normal_style.font.size = Pt(11)

    def _add_text_run(paragraph, text: str, *, bold=False, italic=False, code=False):
        if not text:
            return
        run = paragraph.add_run(text)
        run.bold = bold
        run.italic = italic
        if code:
            run.font.name = "Courier New"
            run.font.size = Pt(10)

    def _render_inline(paragraph, node, *, bold=False, italic=False, code=False):
        tag = (node.tag or "").lower()
        cur_bold = bold or tag in {"strong", "b", "th"}
        cur_italic = italic or tag in {"em", "i"}
        cur_code = code or tag == "code"

        if node.text:
            _add_text_run(
                paragraph,
                node.text,
                bold=cur_bold,
                italic=cur_italic,
                code=cur_code,
            )

        for child in list(node):
            child_tag = (child.tag or "").lower()
            if child_tag == "br":
                _add_text_run(paragraph, "\n", bold=cur_bold, italic=cur_italic, code=cur_code)
            else:
                _render_inline(
                    paragraph,
                    child,
                    bold=cur_bold,
                    italic=cur_italic,
                    code=cur_code,
                )
            if child.tail:
                _add_text_run(
                    paragraph,
                    child.tail,
                    bold=cur_bold,
                    italic=cur_italic,
                    code=cur_code,
                )

    def _plain_text(node) -> str:
        chunks: list[str] = []
        if node.text:
            chunks.append(node.text)
        for c in list(node):
            chunks.append(_plain_text(c))
            if c.tail:
                chunks.append(c.tail)
        return "".join(chunks).strip()

    def _render_table(table_node):
        rows = [r for r in list(table_node) if (r.tag or "").lower() == "tr"]
        if not rows:
            return
        col_count = 0
        for r in rows:
            cells = [c for c in list(r) if (c.tag or "").lower() in {"th", "td"}]
            col_count = max(col_count, len(cells))
        if col_count == 0:
            return

        t = doc.add_table(rows=len(rows), cols=col_count)
        t.style = "Table Grid"
        for i, r in enumerate(rows):
            cells = [c for c in list(r) if (c.tag or "").lower() in {"th", "td"}]
            for j, c in enumerate(cells):
                p = t.cell(i, j).paragraphs[0]
                _render_inline(p, c, bold=(c.tag or "").lower() == "th")

    html_body = _md_to_simple_html(md)
    try:
        root = ET.fromstring(f"<root>{html_body}</root>")
    except ET.ParseError:
        root = ET.fromstring(f"<root><pre>{html.escape(md)}</pre></root>")

    for node in list(root):
        tag = (node.tag or "").lower()
        if tag in {"h1", "h2", "h3", "h4"}:
            level = {"h1": 1, "h2": 2, "h3": 3, "h4": 4}[tag]
            p = doc.add_heading("", level=level)
            _render_inline(p, node, bold=True)
        elif tag == "p":
            p = doc.add_paragraph()
            _render_inline(p, node)
        elif tag in {"ul", "ol"}:
            list_style = "List Bullet" if tag == "ul" else "List Number"
            for li in [x for x in list(node) if (x.tag or "").lower() == "li"]:
                p = doc.add_paragraph(style=list_style)
                _render_inline(p, li)
        elif tag == "blockquote":
            p = doc.add_paragraph(style="Intense Quote")
            _render_inline(p, node)
        elif tag == "pre":
            p = doc.add_paragraph(style="Intense Quote")
            _add_text_run(p, _plain_text(node), code=True)
        elif tag == "table":
            _render_table(node)
        elif tag == "hr":
            doc.add_paragraph("—" * 24)
        else:
            p = doc.add_paragraph()
            _render_inline(p, node)

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _md_to_simple_html(md: str) -> str:
    try:
        import markdown as md_lib

        return md_lib.markdown(md, extensions=["extra", "nl2br", "tables"])
    except Exception:
        return f"<pre>{html.escape(md)}</pre>"


def _build_professional_html(md: str) -> str:
    body = _md_to_simple_html(md)
    css = """
    @page { size: A4; margin: 24mm 18mm 22mm 18mm; }
    body {
      font-family: "Helvetica", "Arial", sans-serif;
      color: #1f2937;
      font-size: 11pt;
      line-height: 1.6;
    }
    h1, h2, h3, h4 {
      color: #111827;
      margin-top: 1.2em;
      margin-bottom: 0.45em;
      line-height: 1.25;
      page-break-after: avoid;
    }
    h1 {
      font-size: 22pt;
      border-bottom: 2px solid #d1d5db;
      padding-bottom: 8px;
      margin-top: 0;
    }
    h2 {
      font-size: 15pt;
      border-left: 4px solid #374151;
      padding-left: 8px;
    }
    h3 { font-size: 12.5pt; }
    h4 { font-size: 11.5pt; color: #374151; }
    p { margin: 0.25em 0 0.8em 0; }
    ul, ol { margin: 0.25em 0 0.8em 1.4em; }
    li { margin: 0.2em 0; }
    code {
      font-family: "Courier New", monospace;
      font-size: 9.5pt;
      background: #f3f4f6;
      border: 1px solid #e5e7eb;
      border-radius: 3px;
      padding: 1px 4px;
    }
    pre {
      white-space: pre-wrap;
      word-wrap: break-word;
      background: #f9fafb;
      border: 1px solid #e5e7eb;
      border-radius: 6px;
      padding: 10px 12px;
      font-family: "Courier New", monospace;
      font-size: 9.5pt;
      line-height: 1.5;
    }
    table {
      width: 100%;
      border-collapse: collapse;
      margin: 0.8em 0 1em 0;
      font-size: 10pt;
    }
    th, td {
      border: 1px solid #d1d5db;
      padding: 6px 8px;
      vertical-align: top;
      text-align: left;
    }
    th {
      background: #f3f4f6;
      font-weight: 700;
    }
    blockquote {
      border-left: 4px solid #9ca3af;
      margin: 0.8em 0;
      padding: 0.3em 0 0.3em 0.8em;
      color: #4b5563;
      background: #f9fafb;
    }
    hr {
      border: 0;
      border-top: 1px solid #d1d5db;
      margin: 1.2em 0;
    }
    """
    return (
        "<!DOCTYPE html><html><head><meta charset=\"utf-8\"/>"
        f"<style>{css}</style></head><body>{body}</body></html>"
    )


def try_markdown_to_pdf_bytes(md: str) -> Optional[bytes]:
    """
    Try xhtml2pdf first (better layout); fall back to fpdf2 plain text.
    Returns None if no backend could produce PDF bytes.
    """
    wrapped = _build_professional_html(md)

    try:
        from xhtml2pdf import pisa

        out = io.BytesIO()
        status = pisa.CreatePDF(src=wrapped, dest=out, encoding="utf-8")
        if status.err:
            raise RuntimeError(f"xhtml2pdf status err={status.err}")
        raw = out.getvalue()
        if raw:
            return raw
    except Exception as e:
        _log.warning("xhtml2pdf backend failed: %s", e, exc_info=True)

    try:
        from fpdf import FPDF
        from fpdf.html import HTML2FPDF

        class _PDF(FPDF):
            def __init__(self) -> None:
                super().__init__()
                self.set_auto_page_break(auto=True, margin=14)

        pdf = _PDF()
        pdf.add_page()
        font_path = Path(__file__).resolve().parent / "fonts" / "NotoSans-Regular.ttf"
        font_bold_path = Path(__file__).resolve().parent / "fonts" / "NotoSans-Bold.ttf"
        if font_path.exists():
            pdf.add_font("NotoSans", style="", fname=str(font_path))
            # HTML parser may request italic / bold-italic variants (<em>/<i>).
            # Map them to available fonts so rendering doesn't fail when italic
            # TTF files are not shipped.
            pdf.add_font("NotoSans", style="I", fname=str(font_path))
            if font_bold_path.exists():
                pdf.add_font("NotoSans", style="B", fname=str(font_bold_path))
                pdf.add_font("NotoSans", style="BI", fname=str(font_bold_path))
            else:
                pdf.add_font("NotoSans", style="B", fname=str(font_path))
                pdf.add_font("NotoSans", style="BI", fname=str(font_path))
            pdf.set_font("NotoSans", size=10)
        else:
            pdf.set_font("Helvetica", size=10)

        html_body = _md_to_simple_html(md)
        inline_css = (
            "<style>"
            "body{font-size:11pt;line-height:1.6;color:#1f2937;}"
            "h1{font-size:22pt;border-bottom:1px solid #d1d5db;padding-bottom:6px;}"
            "h2{font-size:15pt;color:#111827;}"
            "h3{font-size:12pt;color:#111827;}"
            "p{margin:0 0 8px 0;}"
            "ul,ol{margin:0 0 8px 0;}"
            "table{width:100%;border-collapse:collapse;margin:8px 0;}"
            "th,td{border:1px solid #d1d5db;padding:4px 6px;text-align:left;}"
            "th{background:#f3f4f6;}"
            "</style>"
        )
        # fpdf2 HTML writer supports a subset of HTML/CSS; keep styles simple.
        HTML2FPDF_CLASS = HTML2FPDF
        parser = HTML2FPDF_CLASS(pdf, image_map=None)
        parser.feed(inline_css + html_body)
        parser.close()

        out_b = pdf.output(dest="S")
        if isinstance(out_b, (bytes, bytearray)):
            return bytes(out_b)
        if isinstance(out_b, str):
            return out_b.encode("latin-1", errors="replace")
        return bytes(out_b or b"")
    except Exception as e:
        _log.warning("fpdf2 backend failed: %s", e, exc_info=True)
        return None
